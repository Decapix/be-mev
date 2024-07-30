from docx import Document
from docx.shared import Inches
import os
from fpdf import FPDF
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .utils import *
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
import io

def make_pdf(formulaire, qr_path):
    class PDF(FPDF):
        def __init__(self, orientation='P', unit='mm', format='A4'):
            super().__init__(orientation, unit, format)
            # Ajoute la police DejaVu à partir du dossier fonts
            self.add_font('DejaVu', '', 'fonts/DejaVuSans.ttf', uni=True)
            self.add_font('DejaVu', 'B', 'fonts/DejaVuSans-Bold.ttf', uni=True)  # Bold

    pdf = PDF()
    pdf.add_page()
    pdf.set_font("DejaVu", 'B', 14)

    # Calcul de la largeur du texte pour 'nom' et ajout du texte
    nom_width = pdf.get_string_width(formulaire.nom) + 6  # +6 pour une petite marge
    if nom_width + 150 > 190:  # 190 mm est la largeur utilisable typique sur A4
        pdf.multi_cell(150, 10, formulaire.nom, 0, 'L')  # Utilise multi_cell si le texte est trop long
    else:
        pdf.cell(0, 10, formulaire.nom, 0, 1, 'L')  # Utilise cell si le texte est assez court

    # Calcul de la largeur du texte pour 'nom' et ajout du texte
    campagne_width = pdf.get_string_width(formulaire.campagne.nom) + 6  # +6 pour une petite marge
    if campagne_width + 150 > 190:  # 190 mm est la largeur utilisable typique sur A4
        pdf.multi_cell(150, 10, formulaire.campagne.nom, 0, 'L')  # Utilise multi_cell si le texte est trop long
    else:
        pdf.cell(0, 10, formulaire.campagne.nom, 0, 1, 'L')  # Utilise cell si le texte est assez court

    pdf.image(qr_path, x=165, y=20, w=30)
    # Ajouter un texte explicatif sous le QR code
    pdf.set_font("DejaVu", size=8)
    pdf.set_xy(165, 55)  # Ajustez la position y si nécessaire pour que le texte soit juste sous le QR code

    # Utiliser multi_cell pour permettre le retour à la ligne
    text = "Pour répondre au formulaire\nen ligne, scanner"
    pdf.multi_cell(30, 4, text, border=0, align='C')  # La hauteur de la ligne est réglée à 4, ajustez selon le besoin

    # Repositionner pour le reste du texte
    pdf.set_font("DejaVu", size=10)
    pdf.set_y(pdf.get_y() + 2)  # Ajustez cette valeur selon le besoin de l'espacement après le QR code et le texte explicatif

    pdf.multi_cell(190, 5, """Dans le cadre des travaux de rénovation votés lors de la derniere AG, plusieurs aides financieres et solution de financement sont mobilisables nous
vous faisons parvenir ce questionnaire afin d’obtenir les justificatifs nécessaires pour effecuter les demandes de subventions les aides auxquelles
vous pouvez pretendre. Ce questionnaire nous permettra également, le cas echeant, de reprendre contact avec vous au moment du montage des
dossiers. Il est donc très important d’y répondre.""", ln=True)
    pdf.cell(0, 10, f"Précisions :", ln=True)
    pdf.multi_cell(190, 5, "- Ce questionnaire ne constitue pas une demande d’aides.", ln=True)
    pdf.multi_cell(190, 5, "- Ce questionnaire reste entièrement confidentiel à destination unique du bureau d’études. Les données sont exploitées le temps de la mission puis supprimées à la fin de l’étude.", ln=True)
    pdf.multi_cell(190, 5, "- Le retour de ce questionnaire se fait donc uniquement par courrier ou par email à l’adresse en bas de page. (si retour par email : merci d’indiquer le nom de la residence dans l’objet du message)", ln=True)

    if formulaire.identification :
        formulaire.identification.make_pdf(pdf)
    if formulaire.descriptif_du_logement :
        formulaire.descriptif_du_logement.make_pdf(pdf)
    if formulaire.bati :
        formulaire.bati.make_pdf(pdf)
    if formulaire.chauffage_eau_chaude :
        formulaire.chauffage_eau_chaude.make_pdf(pdf)
    if formulaire.ventilation :
        formulaire.ventilation.make_pdf(pdf)
    if formulaire.sondage :
        formulaire.sondage.make_pdf(pdf)
    if formulaire.financement :
        formulaire.financement.make_pdf(pdf)
    if formulaire.situation_professionnelle :
        formulaire.situation_professionnelle.make_pdf(pdf)
    if formulaire.composition_menage :
        formulaire.composition_menage.make_pdf(pdf)
    if formulaire.aides_individuelles :
        formulaire.aides_individuelles.make_pdf(pdf)
    if formulaire.aides_individuelles_question_complementaire :
        formulaire.aides_individuelles_question_complementaire.make_pdf(pdf)
    if formulaire.document_complementaire :
        formulaire.document_complementaire.make_pdf(pdf)

    # Génération du PDF en mémoire
    pdf_buffer = io.BytesIO()
    pdf.output(pdf_buffer, 'F')
    pdf_buffer.seek(0)

    # Sauvegarder le PDF dans S3
    pdf_path = f'{formulaire.id}.pdf'
    default_storage.save(pdf_path, ContentFile(pdf_buffer.read()))

    return pdf_path
    

def make_docx(formulaire, qr_code_path):
    doc = Document()
    
    # Ajouter une table pour positionner le titre et le QR code
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(4.5)
    table.columns[1].width = Inches(1.5)
    
    # Masquer les bordures de la table
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"sz": 0, "color": "#FFFFFF", "val": "single"},
                            bottom={"sz": 0, "color": "#FFFFFF", "val": "single"},
                            start={"sz": 0, "color": "#FFFFFF", "val": "single"},
                            end={"sz": 0, "color": "#FFFFFF", "val": "single"})
    
    # Ajouter le titre du formulaire dans la première colonne
    cell = table.cell(0, 0)
    cell.text = f'{formulaire.nom}'
    cell.text = f'{formulaire.campagne}'
    
    # Ajouter le QR code dans la deuxième colonne
    cell = table.cell(0, 1)
    paragraph = cell.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(qr_code_path, width=Inches(1.0))  # Ajustez la taille au besoin

    # Assurer que le QR code est à droite
    paragraph.alignment = 2  # 2 signifie aligné à droite
    
    # Texte placeholder
    
    doc.add_paragraph("""Dans le cadre des travaux de rénovation votés lors de la derniere AG, plusieurs aides financieres et solution de financement sont mobilisables nous
vous faisons parvenir ce questionnaire afin d’obtenir les justificatifs nécessaires pour effecuter les demandes de subventions les aides auxquelles
vous pouvez pretendre. Ce questionnaire nous permettra également, le cas echeant, de reprendre contact avec vous au moment du montage des
dossiers. Il est donc très important d’y répondre.""")
    doc.add_paragraph("""Précisions :""")
    doc.add_paragraph("""Ce questionnaire ne constitue pas une demande d’aides.""")
    doc.add_paragraph("""Ce questionnaire reste entièrement confidentiel à destination unique du bureau d’études. Les données sont exploitées le temps de la mission
puis supprimées à la fin de l’étude.""")
    doc.add_paragraph("""Le retour de ce questionnaire se fait donc uniquement par courrier ou par email à l’adresse en bas de page. (si retour par email : merci
d’indiquer le nom de la residence dans l’objet du message)""")
    

    # Répétez pour chaque groupe
    
    if formulaire.identification :
        formulaire.identification.make_docx(doc)
    if formulaire.descriptif_du_logement :
        formulaire.descriptif_du_logement.make_docx(doc)
    if formulaire.bati :
        formulaire.bati.make_docx(doc)
    if formulaire.chauffage_eau_chaude :
        formulaire.chauffage_eau_chaude.make_docx(doc)
    if formulaire.ventilation :
        formulaire.ventilation.make_docx(doc)
    if formulaire.sondage :
        formulaire.sondage.make_docx(doc)
    if formulaire.financement :
        formulaire.financement.make_docx(doc)
    if formulaire.situation_professionnelle :
        formulaire.situation_professionnelle.make_docx(doc)
    if formulaire.composition_menage :
        formulaire.composition_menage.make_docx(doc)
    if formulaire.aides_individuelles :
        formulaire.aides_individuelles.make_docx(doc)
    if formulaire.aides_individuelles_question_complementaire :
        formulaire.aides_individuelles_question_complementaire.make_docx(doc)
    if formulaire.document_complementaire :
        formulaire.document_complementaire.make_docx(doc)

   # Sauvegarder le DOCX en mémoire
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    # Sauvegarder le DOCX dans S3
    docx_path = f'{formulaire.id}.docx'
    default_storage.save(docx_path, ContentFile(docx_buffer.read()))

    return docx_path
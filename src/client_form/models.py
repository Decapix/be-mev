from django.db import models
import uuid
from fpdf import FPDF
from constance import config

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Vos modèles existants
# ... (Assurez-vous que vos autres modèles sont définis ici)
CHOICES_VENTILATION = [
    ('OUI', 'Oui'),
    ('NON', 'Non'),
    ('NSP', 'Ne sais pas')
]



class Campagne(models.Model):
    id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    date_de_creation = models.DateTimeField(auto_now_add=True)
    nom = models.CharField(max_length=255)
    description = models.TextField()
    excel = models.FileField(upload_to='media/excels/', blank=True, null=True)
    
    def __str__(self):
        return f'{self.nom}'
        
class Formulaire(models.Model):
    id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    campagne = models.ForeignKey(
        Campagne, on_delete=models.CASCADE, related_name='formulaires', blank=True, null=True)
    nom = models.CharField(max_length=100)
    formulaire_type = models.BooleanField(default=False, blank=True, null=True)
    # Relation Many-to-Many avec différents modèles de groupe
    clone = models.BooleanField(default=False)
    identification = models.OneToOneField('Identification_f',
                                          null=True,
                                          blank=True,
                                          on_delete=models.CASCADE)
    descriptif_du_logement = models.OneToOneField('DescriptifDuLogement_f',
                                                  null=True,
                                                  blank=True,
                                                  on_delete=models.CASCADE)
    bati = models.OneToOneField('BATI_f',
                                null=True,
                                blank=True,
                                on_delete=models.CASCADE)
    chauffage_eau_chaude = models.OneToOneField('ChauffageEauChaude_f',
                                                null=True,
                                                blank=True,
                                                on_delete=models.CASCADE)
    ventilation = models.OneToOneField('Ventilation_f',
                                       null=True,
                                       blank=True,
                                       on_delete=models.CASCADE)
    sondage = models.OneToOneField('Sondage_f',
                                   null=True,
                                   blank=True,
                                   on_delete=models.CASCADE)
    financement = models.OneToOneField('Financement_f',
                                       null=True,
                                       blank=True,
                                       on_delete=models.CASCADE)
    situation_professionnelle = models.OneToOneField(
        'SituationProfessionnelle_fp',
        null=True,
        blank=True,
        on_delete=models.CASCADE)
    composition_menage = models.OneToOneField('CompositionMenage_fp',
                                              null=True,
                                              blank=True,
                                              on_delete=models.CASCADE)

    proprietaires_occupants_intro = models.OneToOneField('ProprietairesOccupantsIntro_fp',
                                                         null=True,
                                                         blank=True,
                                                         on_delete=models.CASCADE)
    aides_individuelles = models.OneToOneField('AidesIndividuelles_fp',
                                               null=True,
                                               blank=True,
                                               on_delete=models.CASCADE)
    aides_individuelles_question_complementaire = models.OneToOneField('AidesIndividuellesQuestionComplementaire_fp',
                                                                       null=True,
                                                                       blank=True,
                                                                       on_delete=models.CASCADE)
    document_complementaire = models.OneToOneField('DocumentComplementaire_f',
                                                   null=True,
                                                   blank=True,
                                                   on_delete=models.CASCADE)

    def __str__(self):
        return f'{self.nom} (ID: {self.id})'

    def get_linked_objects(self):
        linked_objects = []
        fields = [
            self.identification, self.descriptif_du_logement, self.bati,
            self.chauffage_eau_chaude, self.ventilation, self.sondage,
            self.financement, self.situation_professionnelle, self.composition_menage,
            self.aides_individuelles, self.aides_individuelles_question_complementaire,
            self.document_complementaire
        ]
        for obj in fields:
            if obj is not None:
                linked_objects.append(obj)

        return linked_objects


class MiseEnPage(models.Model):
    formulaire = models.OneToOneField(
        'Formulaire',
        on_delete=models.CASCADE,
        related_name='mise_en_page',
        verbose_name='Formulaire associé'
    )
    qr_code = models.ImageField(
        upload_to='media/qr_codes/',
        verbose_name='QR Code',
        blank=True,
        null=True
    )
    pdf = models.FileField(
        upload_to='media/pdfs/',
        verbose_name='PDF',
        blank=True,
        null=True
    )
    docx = models.FileField(
        upload_to='media/docx_files/',
        verbose_name='Document Word',
        blank=True,
        null=True
    )

    def __str__(self):
        return f'Mise en page pour {self.formulaire.nom}'


class Identification_f(models.Model):
    nom = models.CharField(max_length=100)
    telephone = models.CharField(max_length=25)
    prenom = models.CharField(max_length=100)
    email = models.CharField(max_length=100)

    def __str__(self):
        return f'Identification_f {self.id}'

    def make_pdf(self, pdf):
        pdf.set_font("DejaVu", size=10)
        pdf.cell(0, 10, 'Questionnaire d\'Identification', ln=True)
        pdf.cell(0, 10, f"Nom: ___________________________________________", ln=True)
        pdf.cell(0, 10, f"Prénom: ______________________________________", ln=True)
        pdf.cell(0, 10, f"Téléphone: _________________________________", ln=True)
        pdf.cell(0, 10, f"Email: ________________________________________", ln=True)
        pdf.ln(4)  # Ajoute un espace entre les sections

    def make_docx(self, doc):
        doc.add_heading('Questionnaire d\'Identification', level=2)
        doc.add_paragraph(f"Nom: \n__________________________________________")
        doc.add_paragraph(f"Prénom: \n____________________________________")
        doc.add_paragraph(f"Téléphone: \n_______________________________")
        doc.add_paragraph(f"Email:\n______________________________________")
        doc.add_paragraph("")  # Ajoute un espace entre les sections

    def to_excel_row(self):
        return {
            'Nom': self.nom,
            'Téléphone': self.telephone,
            'Prénom': self.prenom,
            'Email': self.email
        }

    def get_description(self):
        # Retourne une description dynamique basée sur les attributs du modèle
        return config.Identification_f_description


class DescriptifDuLogement_f(models.Model):
    numero_du_lot = models.CharField(max_length=50, blank=True)
    proprietaire_occupant = models.BooleanField(
        default=False)  # Utilisation de False comme valeur par défaut
    etage = models.PositiveIntegerField(blank=True, null=True)
    batiment = models.CharField(max_length=100, blank=True, null=True)
    nombre_de_piece = models.PositiveIntegerField(blank=True, null=True)
    surface = models.CharField(max_length=100, blank=True, null=True)
    annee_d_aquisition = models.DateField(blank=True, null=True)

    def __str__(self):
        return f'DescriptifDuLogement_f {self.id}'

    def make_pdf(self, pdf):
        pdf.set_font("DejaVu", size=10)
        pdf.cell(0, 10, 'Questionnaire du Descriptif du Logement', ln=True)
        pdf.cell(0, 10, f"Numéro du lot: __________________________________", ln=True)
        pdf.cell(0, 10, f"Propriétaire occupant: Oui [] Non []", ln=True)
        pdf.cell(0, 10, 'Questionnaire du Logement UN BAT', ln=True)
        pdf.cell(0, 10, f"Étage: _________", ln=True)
        pdf.cell(0, 10, f"Bâtiment: _________", ln=True)
        pdf.cell(0, 10, f"Nombre de pièces: _________", ln=True)
        pdf.cell(0, 10, f"Surface: _________ m²", ln=True)
        pdf.cell(0, 10, f"Année d'acquisition: _________", ln=True)
        pdf.ln(4)

    def make_docx(self, doc):
        doc.add_heading('Questionnaire du Descriptif du Logement', level=2)
        doc.add_paragraph(
            f"Numéro du lot: \n__________________________________")
        doc.add_paragraph(f"Propriétaire occupant: Oui [] Non []")
        doc.add_paragraph("")
        doc.add_heading('Questionnaire du Logement UN BAT', level=2)
        doc.add_paragraph(f"Étage: _________")
        doc.add_paragraph(f"Bâtiment: _________")
        doc.add_paragraph(f"Nombre de pièces: _________")
        doc.add_paragraph(f"Surface: _________ m²")
        doc.add_paragraph(f"Année d'acquisition: _________")
        doc.add_paragraph("")

    def to_excel_row(self):
        return {
            'Numéro du Lot': self.numero_du_lot,
            'Propriétaire Occupant': 'Oui' if self.proprietaire_occupant else 'Non',
            'Étage': self.etage,
            'Bâtiment': self.batiment,
            'Nombre de Pièces': self.nombre_de_piece,
            'Surface (m²)': self.surface,
            'Année d’Acquisition': self.annee_d_aquisition
        }

    def get_description(self):
        # Retourne une description dynamique basée sur les attributs du modèle
        return config.DescriptifDuLogement_f_description


class BATI_f(models.Model):
    TYPE_ISOLANT_CHOICES = [('polystyrene', 'Polystyrène'),
                            ('laine_minerale',
                             'Laine minérale'),
                            ('autre', 'Autre')]
    EPAISSEUR_ISOLANT_CHOICES = [('2_4_cm', '2 à 4 cm'),
                                 ('4_6_cm', '4 à 6 cm'),
                                 ('6_8_cm', '6 à 8 cm'),
                                 ('autre', 'Autre')]
    travaux_engage = models.BooleanField(default=False)  # False = non engage
    type_isolant = models.CharField(max_length=50,
                                    choices=TYPE_ISOLANT_CHOICES,
                                    blank=True,
                                    null=True)
    epaisseur_isolant = models.CharField(max_length=50,
                                         choices=EPAISSEUR_ISOLANT_CHOICES,
                                         blank=True,
                                         null=True)
    autre_isolant = models.CharField(max_length=100, blank=True, null=True)
    autre_epaisseur = models.CharField(max_length=100, blank=True, null=True)

    VITRAGE_CHOICES = [('origne', 'Origne'), ('rénové', 'Rénové')]

    DATE_CHOICES = [('plus_10', '+ de 10 ans'),
                    ('entre_5_10', 'entre 5 et 10 ans'),
                    ('moins_5', '- de 5 ans')]

    VOLET_CHOICES = [('origine', 'Origine'), ('renovee', 'Rénovée')]

    sejour1_vitrage = models.CharField(
        max_length=6, choices=VITRAGE_CHOICES, blank=True, null=True)
    sejour1_date = models.CharField(max_length=10,
                                    choices=DATE_CHOICES,
                                    blank=True,
                                    null=True)
    sejour1_volet = models.CharField(max_length=10,
                                     choices=VOLET_CHOICES,
                                     blank=True,
                                     null=True)

    sejour2_vitrage = models.CharField(
        max_length=6, choices=VITRAGE_CHOICES, blank=True, null=True)
    sejour2_date = models.CharField(max_length=10,
                                    choices=DATE_CHOICES,
                                    blank=True,
                                    null=True)
    sejour2_volet = models.CharField(max_length=10,
                                     choices=VOLET_CHOICES,
                                     blank=True,
                                     null=True)

    cuisine_vitrage = models.CharField(
        max_length=6, choices=VITRAGE_CHOICES, blank=True, null=True)
    cuisine_date = models.CharField(max_length=10,
                                    choices=DATE_CHOICES,
                                    blank=True,
                                    null=True)
    cuisine_volet = models.CharField(max_length=10,
                                     choices=VOLET_CHOICES,
                                     blank=True,
                                     null=True)

    chambre1_vitrage = models.CharField(
        max_length=6, choices=VITRAGE_CHOICES, blank=True, null=True)
    chambre1_date = models.CharField(max_length=10,
                                     choices=DATE_CHOICES,
                                     blank=True,
                                     null=True)
    chambre1_volet = models.CharField(max_length=10,
                                      choices=VOLET_CHOICES,
                                      blank=True,
                                      null=True)

    chambre2_vitrage = models.CharField(
        max_length=6, choices=VITRAGE_CHOICES, blank=True, null=True)
    chambre2_date = models.CharField(max_length=10,
                                     choices=DATE_CHOICES,
                                     blank=True,
                                     null=True)
    chambre2_volet = models.CharField(max_length=10,
                                      choices=VOLET_CHOICES,
                                      blank=True,
                                      null=True)

    chambre3_vitrage = models.CharField(
        max_length=6, choices=VITRAGE_CHOICES, blank=True, null=True)
    chambre3_date = models.CharField(max_length=10,
                                     choices=DATE_CHOICES,
                                     blank=True,
                                     null=True)
    chambre3_volet = models.CharField(max_length=10,
                                      choices=VOLET_CHOICES,
                                      blank=True,
                                      null=True)

    chambre4_vitrage = models.CharField(
        max_length=6, choices=VITRAGE_CHOICES, blank=True, null=True)
    chambre4_date = models.CharField(max_length=10,
                                     choices=DATE_CHOICES,
                                     blank=True,
                                     null=True)
    chambre4_volet = models.CharField(max_length=10,
                                      choices=VOLET_CHOICES,
                                      blank=True,
                                      null=True)

    salle_de_bain_vitrage = models.CharField(max_length=6,
                                             choices=VITRAGE_CHOICES, blank=True, null=True)
    salle_de_bain_date = models.CharField(max_length=10,
                                          choices=DATE_CHOICES,
                                          blank=True,
                                          null=True)
    salle_de_bain_volet = models.CharField(max_length=10,
                                           choices=VOLET_CHOICES,
                                           blank=True,
                                           null=True)

    wc_vitrage = models.CharField(
        max_length=6, choices=VITRAGE_CHOICES, blank=True, null=True)
    wc_date = models.CharField(max_length=10,
                               choices=DATE_CHOICES,
                               blank=True,
                               null=True)
    wc_volet = models.CharField(max_length=10,
                                choices=VOLET_CHOICES,
                                blank=True,
                                null=True)

    def __str__(self):
        return f'BATI {self.id}'

    def make_pdf(self, pdf):
        # Police plus petite pour tout faire tenir
        pdf.set_font("DejaVu", size=10)
        pdf.add_page()

        # Titre du questionnaire
        pdf.cell(0, 10, 'Questionnaire BATI', ln=True, align='C')

        # Questions sur l'isolation
        pdf.cell(
            0, 10, 'Avez-vous engagé des travaux d\'isolation intérieure des murs ?', ln=True)
        pdf.cell(0, 10, 'Oui [ ]   Non [ ]', ln=True)
        pdf.cell(
            0, 10, 'Si oui, quel est le type d\'isolant et l\'épaisseur mise en place ?', ln=True)
        pdf.cell(
            0, 10, 'Polystyrène [ ]   Laine minérale [ ]   Autre (préciser) ___________________', ln=True)
        pdf.cell(
            0, 10, '2 à 4 cm [ ]   4 à 6 cm [ ]   6 à 8 cm [ ]   Autre (préciser) ___________________', ln=True)

        # Espacement
        pdf.ln(4)

        # En-tête pour les détails des pièces
        pdf.set_fill_color(200, 220, 255)  # Couleur de fond pour l'en-tête
        pdf.cell(46, 10, 'Pièce', border=1, ln=0, fill=True)
        pdf.cell(46, 10, 'Type de vitrage', border=1, ln=0, fill=True)
        pdf.cell(46, 10, 'Date du changement', border=1, ln=0, fill=True)
        pdf.cell(46, 10, 'Type de volet', border=1, ln=1, fill=True)

        # Liste des pièces
        pieces = ['Séjour 1', 'Séjour 2', 'Cuisine', 'Chambre 1',
                  'Chambre 2', 'Chambre 3', 'Chambre 4', 'Salle de bain', 'WC']
        for piece in pieces:
            pdf.cell(46, 10, piece, border=1, ln=0)
            pdf.cell(46, 10, 'Origine [ ]  Rénové [ ]', border=1, ln=0)
            pdf.cell(46, 10, 'Date _______________', border=1, ln=0)
            pdf.cell(46, 10, 'Origine [ ]  Rénové [ ]', border=1, ln=1)

    def make_docx(self, doc):
        # Titre du questionnaire
        doc.add_heading('Questionnaire BATI', level=1)

        # Questions sur l'isolation
        doc.add_paragraph(
            'Avez-vous engagé des travaux d\'isolation intérieure des murs ?')
        doc.add_paragraph('Oui [ ]   Non [ ]')
        doc.add_paragraph(
            'Si oui, quel est le type d\'isolant et l\'épaisseur mise en place ?')
        doc.add_paragraph(
            'Polystyrène [ ]   Laine minérale [ ]   Autre (préciser) ___________________')
        doc.add_paragraph(
            '2 à 4 cm [ ]   4 à 6 cm [ ]   6 à 8 cm [ ]   Autre (préciser) ___________________')

        # Espacementzz
        doc.add_paragraph('')

        # Table pour les détails des pièces
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Pièce'
        hdr_cells[1].text = 'Type de vitrage'
        hdr_cells[2].text = 'Date du changement'
        hdr_cells[3].text = 'Type de volet'
        table.style = 'Table Grid'  # Ajout d'un style de grille

        # Liste des pièces
        pieces = ['Séjour 1', 'Séjour 2', 'Cuisine', 'Chambre 1',
                  'Chambre 2', 'Chambre 3', 'Chambre 4', 'Salle de bain', 'WC']
        for piece in pieces:
            row_cells = table.add_row().cells
            row_cells[0].text = piece
            row_cells[1].text = 'Origine []  Rénové [ ]'
            row_cells[2].text = 'Date ________________'
            row_cells[3].text = 'Origine []  Rénové [ ]'

        doc.add_paragraph('')

    def to_excel_row(self):
        return {
            'Travaux Engagés': 'Oui' if self.travaux_engage else 'Non',
            'Type Isolant': self.type_isolant,
            'Épaisseur Isolant': self.epaisseur_isolant,
            'Autre Isolant': self.autre_isolant,
            'Autre Épaisseur': self.autre_epaisseur,
            'Séjour 1 Vitrage': self.sejour1_vitrage,
            'Séjour 1 Date': self.sejour1_date,
            'Séjour 1 Volet': self.sejour1_volet,
            'Séjour 2 Vitrage': self.sejour2_vitrage,
            'Séjour 2 Date': self.sejour2_date,
            'Séjour 2 Volet': self.sejour2_volet,
            'Cuisine Vitrage': self.cuisine_vitrage,
            'Cuisine Date': self.cuisine_date,
            'Cuisine Volet': self.cuisine_volet,
            'Chambre 1 Vitrage': self.chambre1_vitrage,
            'Chambre 1 Date': self.chambre1_date,
            'Chambre 1 Volet': self.chambre1_volet,
            'Chambre 2 Vitrage': self.chambre2_vitrage,
            'Chambre 2 Date': self.chambre2_date,
            'Chambre 2 Volet': self.chambre2_volet,
            'Chambre 3 Vitrage': self.chambre3_vitrage,
            'Chambre 3 Date': self.chambre3_date,
            'Chambre 3 Volet': self.chambre3_volet,
            'Chambre 4 Vitrage': self.chambre4_vitrage,
            'Chambre 4 Date': self.chambre4_date,
            'Chambre 4 Volet': self.chambre4_volet,
            'SDB Vitrage': self.salle_de_bain_vitrage,
            'SDB Date': self.salle_de_bain_date,
            'SDB Volet': self.salle_de_bain_volet,
            'WC Vitrage': self.wc_vitrage,
            'WC Date': self.wc_date,
            'WC Volet': self.wc_volet
        }

    def get_description(self):
        # Retourne une description dynamique basée sur les attributs du modèle
        return config.BATI_f_description


class ChauffageEauChaude_f(models.Model):
    type_chauffage = models.CharField(max_length=50,
                                      choices=[('chaudiere_gaz',
                                                'Chaudière gaz'),
                                               ('electrique', 'Electrique'),
                                               ('autres', 'Autres')])
    chauffage_details = models.TextField(blank=True, null=True)

    type_eau_chaude = models.CharField(max_length=50,
                                       choices=[('chaudiere_gaz',
                                                 'Chaudière gaz'),
                                                ('electrique', 'Electrique'),
                                                ('autres', 'Autres')])
    eau_chaude_details = models.TextField(blank=True, null=True)

    periode_debut = models.DateField(blank=True, null=True)
    periode_fin = models.DateField(blank=True, null=True)
    consommation_kwh = models.FloatField(blank=True, null=True)
    cout_ttc = models.FloatField(blank=True, null=True)

    def __str__(self):
        return f'ChauffageEauChaude {self.id}'

    def make_pdf(self, pdf):
        pdf.set_font("DejaVu", size=10)
        pdf.cell(0, 10, "Détails du Chauffage et de l'Eau Chaude", ln=True)
        pdf.cell(
            0, 10, f"Type de chauffage: ___________________________________", ln=True)
        pdf.cell(
            0, 10, f"Détails du chauffage: ________________________________", ln=True)
        pdf.cell(
            0, 10, f"Type d'eau chaude: __________________________________", ln=True)
        pdf.cell(
            0, 10, f"Détails d'eau chaude: ________________________________", ln=True)
        pdf.cell(
            0, 10, f"Période de début: ___________________________________", ln=True)
        pdf.cell(
            0, 10, f"Période de fin: _____________________________________", ln=True)
        pdf.cell(
            0, 10, f"Consommation en kWh: ________________________________", ln=True)
        pdf.cell(
            0, 10, f"Coût total TTC: _____________________________________", ln=True)
        pdf.ln(4)

    def make_docx(self, doc):
        doc.add_heading("Détails du Chauffage et de l'Eau Chaude", level=2)
        doc.add_paragraph(
            f"Type de chauffage: \n__________________________________")
        doc.add_paragraph(
            f"Détails du chauffage: \n________________________________")
        doc.add_paragraph(
            f"Type d'eau chaude: \n__________________________________")
        doc.add_paragraph(
            f"Détails d'eau chaude: \n________________________________")
        doc.add_paragraph(
            f"Période de début: \n__________________________________")
        doc.add_paragraph(
            f"Période de fin: \n____________________________________")
        doc.add_paragraph(
            f"Consommation en kWh: \n______________________________")
        doc.add_paragraph(
            f"Coût total TTC: \n____________________________________")
        doc.add_paragraph("")  # Ajoute un espace entre les sections

    def to_excel_row(self):
        return {
            'Type de Chauffage': self.type_chauffage,
            'Détails Chauffage': self.chauffage_details,
            'Type d\'Eau Chaude': self.type_eau_chaude,
            'Détails Eau Chaude': self.eau_chaude_details,
            'Début de la Période': self.periode_debut,
            'Fin de la Période': self.periode_fin,
            'Consommation kWh': self.consommation_kwh,
            'Coût TTC': self.cout_ttc
        }

    def get_description(self):
        # Retourne une description dynamique basée sur les attributs du modèle
        return config.ChauffageEauChaude_f_description


class Ventilation_f(models.Model):

    grilles_entree_air = models.CharField(
        max_length=3, choices=CHOICES_VENTILATION, null=True)
    bouches_extraction_air = models.CharField(
        max_length=3, choices=CHOICES_VENTILATION, null=True)
    nettoyage_regulier = models.CharField(
        max_length=3, choices=CHOICES_VENTILATION, null=True)
    ventilation_motorisee = models.CharField(
        max_length=3, choices=CHOICES_VENTILATION, null=True)
    ventilation_ouverte_temps = models.CharField(
        max_length=100, blank=True, null=True)

    def __str__(self):
        return f'Ventilation {self.id}'

    def make_pdf(self, pdf):
        pdf.set_font("DejaVu", size=10)
        pdf.cell(0, 10, 'Questionnaire de Ventilation', ln=True)
        pdf.cell(
            0, 10, f"Grilles d'entrée d'air: Oui [] Non [] Je ne sais pas []", ln=True)
        pdf.cell(
            0, 10, f"Bouches d'extraction d'air: Oui [] Non [] Je ne sais pas []", ln=True)
        pdf.cell(
            0, 10, f"Nettoyage régulier: Oui [] Non [] Je ne sais pas []", ln=True)
        pdf.cell(
            0, 10, f"Ventilation motorisée: Oui [] Non [] Je ne sais pas []", ln=True)
        pdf.cell(0, 10, f"Temps ouverte pour ventilation: {self.ventilation_ouverte_temps}", ln=True)
        pdf.ln(4)  # Ajoute un espace entre les sections

    def make_docx(self, doc):
        doc.add_heading('Questionnaire de Ventilation', level=2)
        doc.add_paragraph(
            f"Grilles d'entrée d'air: Oui [] Non [] Je ne sais pas []")
        doc.add_paragraph(
            f"Bouches d'extraction d'air: Oui [] Non [] Je ne sais pas []")
        doc.add_paragraph(
            f"Nettoyage régulier: Oui [] Non [] Je ne sais pas []")
        doc.add_paragraph(
            f"Ventilation motorisée: Oui [] Non [] Je ne sais pas []")
        doc.add_paragraph(f"Temps ouverte pour ventilation: {self.ventilation_ouverte_temps}")
        doc.add_paragraph("")  # Ajoute un espace entre les sections

    def to_excel_row(self):
        return {
            'Grilles Entrée Air': self.grilles_entree_air,
            'Bouches Extraction Air': self.bouches_extraction_air,
            'Nettoyage Régulier': self.nettoyage_regulier,
            'Ventilation Motorisée': self.ventilation_motorisee,
            'Ventilation Ouverte Temps': self.ventilation_ouverte_temps
        }

    def get_description(self):
        # Retourne une description dynamique basée sur les attributs du modèle
        return config.Ventilation_f_description


class Sondage_f(models.Model):
    isolation_facades = models.BooleanField(default=False)
    isolation_toiture = models.BooleanField(default=False)
    regulation_chauffage = models.BooleanField(default=False)
    remplacement_fenetres = models.BooleanField(default=False)
    amelioration_ventilation = models.BooleanField(default=False)
    remplacement_chauffage = models.BooleanField(default=False)

    def __str__(self):
        return f'Sondage {self.id}'

    def make_pdf(self, pdf):
        pdf.set_font("DejaVu", size=10)
        pdf.cell(0, 10, 'Sondage des améliorations de maison', ln=True)
        pdf.cell(0, 10, f"Isolation des façades: Oui [] Non []", ln=True)
        pdf.cell(0, 10, f"Isolation de la toiture: Oui [] Non []", ln=True)
        pdf.cell(0, 10, f"Régulation du chauffage: Oui [] Non []", ln=True)
        pdf.cell(0, 10, f"Remplacement des fenêtres: Oui [] Non []", ln=True)
        pdf.cell(0, 10, f"Amélioration de la ventilation: Oui [] Non []", ln=True)
        pdf.cell(0, 10, f"Remplacement du chauffage: Oui [] Non []", ln=True)
        pdf.ln(4)  # Ajoute un espace entre les sections

    def make_docx(self, doc):
        doc.add_heading('Sondage des améliorations de maison', level=2)
        doc.add_paragraph(f"Isolation des façades: Oui [] Non []")
        doc.add_paragraph(f"Isolation de la toiture: Oui [] Non []")
        doc.add_paragraph(f"Régulation du chauffage: Oui [] Non []")
        doc.add_paragraph(f"Remplacement des fenêtres: Oui [] Non []")
        doc.add_paragraph(f"Amélioration de la ventilation: Oui [] Non []")
        doc.add_paragraph(f"Remplacement du chauffage: Oui [] Non []")
        doc.add_paragraph("")  # Ajoute un espace entre les sections

    def to_excel_row(self):
        return {
            'Isolation des Façades': 'Oui' if self.isolation_facades else 'Non',
            'Isolation de la Toiture': 'Oui' if self.isolation_toiture else 'Non',
            'Régulation du Chauffage': 'Oui' if self.regulation_chauffage else 'Non',
            'Remplacement des Fenêtres': 'Oui' if self.remplacement_fenetres else 'Non',
            'Amélioration de la Ventilation': 'Oui' if self.amelioration_ventilation else 'Non',
            'Remplacement du Chauffage': 'Oui' if self.remplacement_chauffage else 'Non'
        }

    def get_description(self):
        # Retourne une description dynamique basée sur les attributs du modèle
        return config.Sondage_f_description


class Financement_f(models.Model):
    pret_collectif = models.BooleanField(default=False)
    pret_individuel = models.BooleanField(default=False)
    financement_fonds_propres = models.BooleanField(default=False)
    ne_se_prononce_pas = models.BooleanField(default=False)

    duree_pret = models.IntegerField(choices=[(3, '3 ans'), (5, '5 ans'),
                                              (7, '7 ans'), (10, '10 ans'),
                                              (12, '12 ans'), (15, '15 ans'),
                                              (20, '20 ans')],
                                     blank=True,
                                     null=True)

    def __str__(self):
        return f'Financement {self.id}'

    def make_pdf(self, pdf):
        pdf.set_font("DejaVu", size=10)
        pdf.cell(0, 10, 'Questionnaire de Financement', ln=True)
        pdf.cell(0, 10, f"Prêt collectif: Oui [] Non []", ln=True)
        pdf.cell(0, 10, f"Prêt individuel: Oui [] Non []", ln=True)
        pdf.cell(0, 10, f"Financement par fonds propres: Oui [] Non []", ln=True)
        pdf.cell(0, 10, f"Ne se prononce pas: Oui [] Non []", ln=True)
        pdf.cell(0, 10, f"Durée du prêt: ________ ans", ln=True)
        pdf.ln(4)  # Ajoute un espace entre les sections

    def make_docx(self, doc):
        doc.add_heading('Questionnaire de Financement', level=2)
        doc.add_paragraph(f"Prêt collectif: Oui [] Non []")
        doc.add_paragraph(f"Prêt individuel: Oui [] Non []")
        doc.add_paragraph(f"Financement par fonds propres: Oui [] Non []")
        doc.add_paragraph(f"Ne se prononce pas: Oui [] Non []")
        doc.add_paragraph(f"Durée du prêt: ________ ans")
        doc.add_paragraph("")  # Ajoute un espace entre les sections

    def to_excel_row(self):
        # Convert boolean to 'Oui'/'Non' for better readability in Excel
        return {
            'Prêt Collectif': 'Oui' if self.pret_collectif else 'Non',
            'Prêt Individuel': 'Oui' if self.pret_individuel else 'Non',
            'Financement par Fonds Propres': 'Oui' if self.financement_fonds_propres else 'Non',
            'Ne se prononce pas': 'Oui' if self.ne_se_prononce_pas else 'Non',
            'Durée du Prêt (ans)': self.get_duree_pret_display() if self.duree_pret else 'N/A'
        }

    def get_description(self):
        # Retourne une description dynamique basée sur les attributs du modèle
        return config.Financement_f_description
# reserver au proprietaire occupant


class SituationProfessionnelle_fp(models.Model):
    SITUATION_CHOICES = [('salarie_prive', 'Salarié du secteur privé'),
                         ('travailleur_independant',
                          'Travailleur indépendant'), ('retraite', 'Retraité'),
                         ('fonctionnaire', 'Fonctionnaire'),
                         ('sans_emploi', 'Sans emploi'), ('autre', 'Autre')]

    situation_professionnelle = models.CharField(max_length=50,
                                                 choices=SITUATION_CHOICES)
    fonctionnaire_details = models.CharField(max_length=100,
                                             blank=True,
                                             null=True)
    situation_professionnelle_conjoint = models.CharField(
        max_length=50, choices=SITUATION_CHOICES)
    fonctionnaire_conjoint_details = models.CharField(max_length=100,
                                                      blank=True,
                                                      null=True)

    beneficie_prestation_caf = models.BooleanField(
        default=False)  # True = Oui, False = Non
    prestation_apa = models.BooleanField(default=False)
    prestation_pch = models.BooleanField(default=False)
    prestation_actp = models.BooleanField(default=False)
    prestation_psd = models.BooleanField(default=False)

    def __str__(self):
        return f'SituationProfessionnelle {self.id}'

    def make_pdf(self, pdf):
        pdf.set_font("DejaVu", size=10)
        pdf.cell(0, 10, 'Questionnaire de Situation Professionnelle', ln=True)
        pdf.cell(0, 10, f"Situation professionnelle: ______________________", ln=True)
        pdf.cell(0, 10, f"Détails fonctionnaire: _________________________", ln=True)
        pdf.cell(
            0, 10, f"Situation professionnelle conjoint: _______________", ln=True)
        pdf.cell(
            0, 10, f"Détails fonctionnaire conjoint: __________________", ln=True)
        pdf.cell(0, 10, f"Bénéficie de prestations CAF: Oui [] Non []", ln=True)
        pdf.cell(0, 10, f"Bénéficie de prestations APA: Oui [] Non []", ln=True)
        pdf.cell(0, 10, f"Bénéficie de prestations PCH: Oui [] Non []", ln=True)
        pdf.cell(0, 10, f"Bénéficie de prestations ACTP: Oui [] Non []", ln=True)
        pdf.cell(0, 10, f"Bénéficie de prestations PSD: Oui [] Non []", ln=True)
        pdf.ln(4)  # Ajoute un espace entre les sections

    def make_docx(self, doc):
        doc.add_heading('Questionnaire de Situation Professionnelle', level=2)
        doc.add_paragraph(
            f"Situation professionnelle:\n______________________")
        doc.add_paragraph(f"Détails fonctionnaire:\n_______________________")
        doc.add_paragraph(
            f"Situation professionnelle conjoint:\n_____________")
        doc.add_paragraph(f"Détails fonctionnaire conjoint:\n________________")
        doc.add_paragraph(f"Bénéficie de prestations CAF: Oui [] Non []")
        doc.add_paragraph(f"Bénéficie de prestations APA: Oui [] Non []")
        doc.add_paragraph(f"Bénéficie de prestations PCH: Oui [] Non []")
        doc.add_paragraph(f"Bénéficie de prestations ACTP: Oui [] Non []")
        doc.add_paragraph(f"Bénéficie de prestations PSD: Oui [] Non []")
        doc.add_paragraph("")  # Ajoute un espace entre les sections

    def to_excel_row(self):
        return {
            'Situation Professionnelle': self.get_situation_professionnelle_display(),
            'Détails Fonctionnaire': self.fonctionnaire_details,
            'Situation Professionnelle Conjoint': self.get_situation_professionnelle_conjoint_display(),
            'Détails Fonctionnaire Conjoint': self.fonctionnaire_conjoint_details,
            'Bénéficie Prestation CAF': 'Oui' if self.beneficie_prestation_caf else 'Non',
            'Prestation APA': 'Oui' if self.prestation_apa else 'Non',
            'Prestation PCH': 'Oui' if self.prestation_pch else 'Non',
            'Prestation ACTP': 'Oui' if self.prestation_actp else 'Non',
            'Prestation PSD': 'Oui' if self.prestation_psd else 'Non'
        }

    def get_description(self):
        # Retourne une description dynamique basée sur les attributs du modèle
        return config.SituationProfessionnelle_fp_description


class CompositionMenage_fp(models.Model):
    SITUATION_CHOICES = [('salarie', 'Salarié'),
                         ('liberal', 'Libéral, indépendant, autoentrepreneur'),
                         ('retraite', 'Retraité'),
                         ('demandeur_emploi', 'Demandeur d\'emploi'),
                         ('etudiant', 'Etudiant, en formation'),
                         ('autre', 'Autre')]

    situation = models.CharField(max_length=50, choices=SITUATION_CHOICES)
    situation_details = models.CharField(max_length=100, blank=True, null=True)

    nombre_personnes = models.IntegerField(null=True)
    nombre_adultes = models.IntegerField(null=True)
    nombre_enfants_mineurs = models.IntegerField(null=True)
    nombre_enfants_majeurs = models.IntegerField(null=True)

    personne_handicap = models.BooleanField(
        default=False)  # True = Oui, False = Non

    def __str__(self):
        return f'CompositionMenage {self.id}'

    def make_pdf(self, pdf):
        pdf.set_font("DejaVu", size=10)
        pdf.cell(0, 10, 'Composition du ménage', ln=True)
        for choice, label in self.SITUATION_CHOICES:
            checkbox = '[ ]' if getattr(self, 'situation') != choice else '[X]'
            pdf.cell(0, 10, f"{checkbox} {label}", ln=True)
        pdf.cell(0, 10, f"Autre (préciser): {self.situation_details}", ln=True)
        pdf.ln(4)  # Ajoute un espace entre les sections

    def make_docx(self, doc):
        doc.add_heading('Composition du ménage', level=2)
        for choice, label in self.SITUATION_CHOICES:
            checkbox = '[ ]' if getattr(self, 'situation') != choice else '[X]'
            doc.add_paragraph(f"{checkbox} {label}")
        doc.add_paragraph(f"Autre (préciser): {self.situation_details}")
        doc.add_paragraph("")  # Ajoute un espace entre les sections

    def to_excel_row(self):
        return {
            'Situation du Ménage': self.get_situation_display(),
            'Détails de la Situation': self.situation_details,
            'Nombre Total de Personnes': self.nombre_personnes,
            'Nombre d\'Adultes': self.nombre_adultes,
            'Nombre d\'Enfants Mineurs': self.nombre_enfants_mineurs,
            'Nombre d\'Enfants Majeurs': self.nombre_enfants_majeurs,
            'Personne(s) en Situation de Handicap': 'Oui' if self.personne_handicap else 'Non'
        }

    def get_description(self):
        # Retourne une description dynamique basée sur les attributs du modèle
        return config.CompositionMenage_fp_description


class ProprietairesOccupantsIntro_fp(models.Model):
    residence_principale = models.BooleanField(
        default=False)  # True = Oui, False = Non
    difficulte_payer_charges = models.BooleanField(
        default=False)  # True = Oui, False = Non
    montant_impayes = models.FloatField(blank=True, null=True)

    def __str__(self):
        return f'ProprietairesOccupants {self.id}'

    def make_pdf(self, pdf):
        pdf.set_font("DejaVu", size=10)
        pdf.cell(0, 10, 'Introducion propriétaire occupant', ln=True)
        pdf.cell(0, 10, f"Résidence principale: Oui [] Non []", ln=True)
        pdf.cell(0, 10, f"Difficulté à payer les charges: Oui [] Non []", ln=True)
        pdf.cell(0, 10, f"Montant impayé: _______ ", ln=True)

    def make_docx(self, doc):
        doc.add_heading('Introducion propriétaire occupant', level=2)
        doc.add_paragraph(f"Résidence principale: Oui [] Non []")
        doc.add_paragraph(f"Difficulté à payer les charges: Oui [] Non []")
        doc.add_paragraph(f"Montant impayé: _______ ")
        doc.add_paragraph("")  # Ajoute un espace entre les sections

    def to_excel_row(self):
        return {
            'Résidence Principale': 'Oui' if self.residence_principale else 'Non',
            'Difficulté à Payer les Charges': 'Oui' if self.difficulte_payer_charges else 'Non',
            'Montant des Impayés': self.montant_impayes if self.montant_impayes is not None else 'N/A'
        }

    def get_description(self):
        # Retourne une description dynamique basée sur les attributs du modèle
        return config.ProprietairesOccupantsIntro_fp_description


def set_cell_background(cell, color):
    cell_properties = cell._element.get_or_add_tcPr()
    cell_shading = OxmlElement('w:shd')
    cell_shading.set(qn('w:fill'), color)
    cell_properties.append(cell_shading)


class AidesIndividuelles_fp(models.Model):
    PROFIL_CHOICES = [('Bleu', 'Bleu'),
                      ('Jaune', 'Jaune'),
                      ('Violet', 'Violet'),
                      ('Rose', 'Rose')]

    profile_menage = models.CharField(max_length=50, choices=PROFIL_CHOICES)

    def __str__(self):
        return f'AidesIndividuelles {self.id}'

    def make_pdf(self, pdf):
        # Police plus petite pour tout faire tenir
        pdf.set_font("DejaVu", size=8)
        pdf.add_page()

        # Titre du questionnaire
        pdf.cell(0, 10, 'AIDES INDIVIDUELLES - Formulaire d\'Evaluation',
                 ln=True, align='C')

        # Sous-titre et instructions
        pdf.cell(0, 10, 'Vous pouvez bénéficier d\'un bonus aux aides si votre revenu fiscal de référence est inférieur aux plafonds suivants:', ln=True)
        pdf.cell(
            0, 10, '(Pour vous situer, merci de vous référer à votre dernier avis d\'imposition.)', ln=True)
        pdf.cell(0, 10, 'Précision: Si plusieurs déclarations de revenus composent le ménage, additionner vos revenus fiscaux de référence.', ln=True)

        # Tableau des plafonds de revenus
        col_widths = [35, 35, 35, 35, 35]  # Largeurs des colonnes
        row_height = 10  # Hauteur des lignes

        # En-tête du tableau
        headers = ['Nombre de personnes', 'Ménage Bleu',
                   'Ménage Jaune', 'Ménage Violet', 'Ménage Rose']
        colors = [(200, 220, 255), (173, 216, 230), (255, 255, 102),
                  (230, 230, 250), (255, 182, 193)]

        for header, color in zip(headers, colors):
            pdf.set_fill_color(*color)
            pdf.cell(col_widths[headers.index(header)], row_height,
                     header, border=1, fill=True, ln=0, align='C')
        pdf.ln()

        # Lignes de données
        for i in range(1, 6):  # pour 1 à 5 personnes
            pdf.cell(col_widths[0], row_height, str(
                i), border=1, ln=0, align='C')
            for j, profil in enumerate(['Bleu', 'Jaune', 'Violet', 'Rose'], start=1):
                key = f'aide_individuel-{profil}{i}'
                value = getattr(config, key, 'default_value')
                pdf.cell(col_widths[j], row_height,
                         value, border=1, ln=0, align='C')
            pdf.ln()

        # Ajouter une ligne pour chaque personne supplémentaire
        pdf.cell(col_widths[0], row_height,
                 'Par personne\nsupplémentaire', border=1, ln=0, align='C')
        for j, profil in enumerate(['Bleu', 'Jaune', 'Violet', 'Rose'], start=1):
            key = f'aide_individuel-{profil}+'
            value = getattr(config, key, 'default_value')
            pdf.cell(col_widths[j], row_height,
                     value, border=1, ln=0, align='C')
        pdf.ln()

        # Ajouter une ligne avec des cases à cocher
        pdf.cell(col_widths[0], row_height,
                 'Sélectionner', border=1, ln=0, align='C')
        for _ in range(4):
            pdf.cell(col_widths[1], row_height,
                     '[ ]', border=1, ln=0, align='C')
        pdf.ln(4)

    def make_docx(self, doc):
        # Ajout d'un titre au document
        title = doc.add_heading(level=1)
        title_run = title.add_run(
            'AIDES INDIVIDUELLES - Formulaire d\'Evaluation')
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Introduction et instructions
        doc.add_paragraph(
            'Vous pouvez bénéficier d\'un bonus aux aides si votre revenu fiscal de référence est inférieur aux plafonds suivants:')
        doc.add_paragraph(
            '(Pour vous situer, merci de vous référer à votre dernier avis d\'imposition.)')
        doc.add_paragraph(
            'Précision: Si plusieurs déclarations de revenus composent le ménage, additionner vos revenus fiscaux de référence.')

        # Ajout d'un tableau pour les plafonds de revenu
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        headers = ['Nombre de personnes', 'Ménage Bleu',
                   'Ménage Jaune', 'Ménage Violet', 'Ménage Rose']
        colors = ['ADD8E6', 'FFFF66', 'E6E6FA', 'FFB6C1']

        for i, (hdr_cell, color) in enumerate(zip(hdr_cells, ['C8DCF0'] + colors)):
            hdr_cell.text = headers[i]
            set_cell_background(hdr_cell, color)

        # Remplissage du tableau avec des données
        for i in range(1, 6):  # pour 1 à 5 personnes
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            for j, profil in enumerate(['Bleu', 'Jaune', 'Violet', 'Rose']):
                key = f'aide_individuel-{profil}{i}'
                value = getattr(config, key, 'default_value')
                row_cells[j+1].text = value

        # Ajouter une ligne pour chaque personne supplémentaire
        row_cells = table.add_row().cells
        row_cells[0].text = 'Par personne supplémentaire'
        for j, profil in enumerate(['Bleu', 'Jaune', 'Violet', 'Rose']):
            key = f'aide_individuel-{profil}+'
            value = getattr(config, key, 'default_value')
            row_cells[j+1].text = value

        # Ajouter une ligne avec des cases à cocher
        row_cells = table.add_row().cells
        row_cells[0].text = 'Sélectionner'
        for i in range(1, 5):
            row_cells[i].text = '[ ]'

        # Ajout d'espacement après le tableau
        doc.add_paragraph('')

    def to_excel_row(self):
        return {
            'Profil du Ménage': self.get_profile_menage_display()
        }

    def get_description(self):
        # Retourne une description dynamique basée sur les attributs du modèle
        return config.AidesIndividuelles_fp_description

    def get_html(self):
        # Sous-titre et instructions
        html = """
        <div class="container">
            <p>(Pour vous situer, merci de vous référer à votre dernier avis d'imposition.)</p>
            <p>Précision: Si plusieurs déclarations de revenus composent le ménage, additionner vos revenus fiscaux de référence.</p>
            
            <table class="table table-bordered">
                <thead>
                    <tr style="background-color: #C8DCF0;">
                        <th>Nombre de personnes</th>
                        <th style="background-color: #ADD8E6;">Ménage Bleu</th>
                        <th style="background-color: #FFFF66;">Ménage Jaune</th>
                        <th style="background-color: #E6E6FA;">Ménage Violet</th>
                        <th style="background-color: #FFB6C1;">Ménage Rose</th>
                    </tr>
                </thead>
                <tbody>"""

        # Lignes de données pour 1 à 5 personnes
        colors = ['#ADD8E6', '#FFFF66', '#E6E6FA', '#FFB6C1']
        for i in range(1, 6):
            html += f"<tr><td align='center'>{i}</td>"
            for j, profil in enumerate(['Bleu', 'Jaune', 'Violet', 'Rose']):
                key = f'aide_individuel-{profil}{i}'
                value = getattr(config, key, 'default_value')  # Assurez-vous que config est défini correctement
                html += f"<td style='background-color: {colors[j]};' align='center'>{value}</td>"
            html += "</tr>"

        # Ligne pour chaque personne supplémentaire
        html += "<tr><td align='center'>Par personne supplémentaire</td>"
        for j, profil in enumerate(['Bleu', 'Jaune', 'Violet', 'Rose']):
            key = f'aide_individuel-{profil}+'
            value = getattr(config, key, 'default_value')  # Utilisez la même logique pour récupérer les valeurs
            html += f"<td style='background-color: {colors[j]};' align='center'>{value}</td>"
        html += "</tr></tbody></table></div>"
        
        return html


class AidesIndividuellesQuestionComplementaire_fp(models.Model):

    revenu_fiscal_foyer = models.FloatField(blank=True, null=True)
    impot_revenu = models.FloatField(blank=True, null=True)
    pret_taux_zero = models.BooleanField(
        default=False)  # True = Oui, False = Non
    aide_anah = models.BooleanField(default=False)  # True = Oui, False = Non
    montant_aide_anah = models.FloatField(blank=True, null=True)
    annee_aide_anah = models.IntegerField(blank=True, null=True)

    def __str__(self):
        return f'AidesIndividuellesQuestionComplementaire {self.id}'

    def make_pdf(self, pdf):
        # Police plus petite pour tout faire tenir
        pdf.set_font("DejaVu", size=8)
        pdf.add_page()

        # Section des informations fiscales et aides
        pdf.cell(0, 10, 'Informations Fiscales et Aides Reçues:', ln=True)
        pdf.cell(
            0, 10, 'Revenu fiscal du foyer: _____________________________ ', ln=True)
        pdf.cell(
            0, 10, 'Montant de l\'impôt sur le revenu: ____________________ ', ln=True)
        pdf.cell(
            0, 10, 'Avez-vous contracté un prêt à taux zéro de l\'État pour l\'achat de votre logement dans les 5 dernières années ? Oui [ ] Non [ ]', ln=True)
        pdf.cell(
            0, 10, 'Avez-vous bénéficié d\'une aide ANAH pour le logement dans les 5 dernières années ? Oui [ ] Non [ ], préciser le montant et l\'année : ____________  en ______', ln=True)

    def make_docx(self, doc):
        # Ajout d'un titre au document

        # Section des informations fiscales et aides
        doc.add_paragraph('Informations Fiscales et Aides Reçues:')
        doc.add_paragraph(
            'Revenu fiscal du foyer: _____________________________ ')
        doc.add_paragraph(
            'Montant de l\'impôt sur le revenu: ____________________ ')
        doc.add_paragraph(
            'Avez-vous contracté un prêt à taux zéro de l\'État pour l\'achat de votre logement dans les 5 dernières années ? Oui [ ] Non [ ]')
        doc.add_paragraph(
            'Avez-vous bénéficié d\'une aide ANAH pour le logement dans les 5 dernières années ? Oui [ ] Non [ ], préciser le montant et l\'année : ____________  en ______')

    def to_excel_row(self):
        return {
            'Revenu Fiscal de Référence du Foyer': self.revenu_fiscal_foyer if self.revenu_fiscal_foyer is not None else 'N/A',
            'Impôt sur le Revenu': self.impot_revenu if self.impot_revenu is not None else 'N/A',
            'Prêt à Taux Zéro': 'Oui' if self.pret_taux_zero else 'Non',
            'Aide de l\'ANAH': 'Oui' if self.aide_anah else 'Non',
            'Montant de l\'Aide ANAH': self.montant_aide_anah if self.montant_aide_anah is not None else 'N/A',
            'Année de l\'Aide ANAH': self.annee_aide_anah if self.annee_aide_anah is not None else 'N/A'
        }

    def get_description(self):
        # Retourne une description dynamique basée sur les attributs du modèle
        return config.AidesIndividuellesQuestionComplementaire_fp_description


class DocumentComplementaire_f(models.Model):
    doc1 = models.FileField(upload_to='media/doc/', blank=True, null=True)
    doc2 = models.FileField(upload_to='media/doc/', blank=True, null=True)
    doc3 = models.FileField(upload_to='media/doc/', blank=True, null=True)
    doc4 = models.FileField(upload_to='media/doc/', blank=True, null=True)
    doc5 = models.FileField(upload_to='media/doc/', blank=True, null=True)

    def __str__(self):
        return f'DocumentComplementaire {self.id}'

    def make_pdf(self, pdf):
        # Police plus petite pour tout faire tenir
        pdf.set_font("DejaVu", size=8)
        pdf.add_page()
        # Section des informations fiscales et aides
        pdf.cell(0, 10, 'Documents complémentaire', ln=True)

    def make_docx(self, doc):
        doc.add_paragraph('Documents complémentaire')

    def to_excel_row(self):
        # Création d'un dictionnaire pour stocker les noms de fichier
        docs = [self.doc1, self.doc2, self.doc3, self.doc4, self.doc5]
        return {f'Document {i+1}': doc.name if doc else '' for i, doc in enumerate(docs)}

    def get_description(self):
        # Retourne une description dynamique basée sur les attributs du modèle
        return config.DocumentComplementaire_f_description


related_fields = {
    'identification': Identification_f,
    'descriptif_du_logement': DescriptifDuLogement_f,
    'bati': BATI_f,
    'chauffage_eau_chaude': ChauffageEauChaude_f,
    'ventilation': Ventilation_f,
    'sondage': Sondage_f,
    'financement': Financement_f,
    'situation_professionnelle': SituationProfessionnelle_fp,
    'composition_menage': CompositionMenage_fp,
    'proprietaires_occupants_intro': ProprietairesOccupantsIntro_fp,
    'aides_individuelles': AidesIndividuelles_fp,
    'aides_individuelles_question_complementaire': AidesIndividuellesQuestionComplementaire_fp,
    "document_complementaire": DocumentComplementaire_f
}
